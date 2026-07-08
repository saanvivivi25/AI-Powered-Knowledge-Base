import asyncio
import datetime
import enum
import os
import re
import time
import uuid
from dataclasses import dataclass
from io import BytesIO
from logging.handlers import RotatingFileHandler
import logging

import bcrypt
from docx import Document as DocxDocument
from fastapi import (
    BackgroundTasks,
    Depends,
    FastAPI,
    HTTPException,
    Request,
    UploadFile,
    status,
)
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.security import OAuth2PasswordBearer
from jose import JWTError, jwt
from langchain_core.prompts import PromptTemplate
from pydantic import BaseModel, ConfigDict, EmailStr, Field
from pydantic_settings import BaseSettings, SettingsConfigDict
from pypdf import PdfReader
from sqlalchemy import DateTime, Enum, ForeignKey, Integer, String, Text, or_, select
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine
from sqlalchemy.orm import DeclarativeBase, Mapped, mapped_column, relationship


# ========================================================================
# CONFIG
# ========================================================================

class Settings(BaseSettings):
    # Auth
    SECRET_KEY: str = "change-this-to-a-long-random-secret"
    ALGORITHM: str = "HS256"
    ACCESS_TOKEN_EXPIRE_MINUTES: int = 60

    # Database
    DATABASE_URL: str = "sqlite+aiosqlite:///./knowledge_base.db"

    # Gemini / AI
    GEMINI_API_KEY: str = ""
    GEMINI_MODEL: str = "gemini-1.5-flash"

    # Storage
    UPLOAD_DIR: str = "uploads"
    LOG_DIR: str = "logs"
    MAX_UPLOAD_MB: int = 20

    model_config = SettingsConfigDict(env_file=".env", extra="ignore")


settings = Settings()


# ========================================================================
# LOGGING (Module 6) -> writes to logs/app.log
# ========================================================================

def _build_logger() -> logging.Logger:
    os.makedirs(settings.LOG_DIR, exist_ok=True)
    log_path = os.path.join(settings.LOG_DIR, "app.log")

    _logger = logging.getLogger("kb_api")
    _logger.setLevel(logging.INFO)

    formatter = logging.Formatter("%(asctime)s | %(levelname)-8s | %(name)s | %(message)s")

    file_handler = RotatingFileHandler(log_path, maxBytes=5 * 1024 * 1024, backupCount=3, encoding="utf-8")
    file_handler.setFormatter(formatter)
    file_handler.setLevel(logging.INFO)

    console_handler = logging.StreamHandler()
    console_handler.setFormatter(formatter)
    console_handler.setLevel(logging.INFO)

    _logger.addHandler(file_handler)
    _logger.addHandler(console_handler)
    _logger.propagate = False
    return _logger


logger = _build_logger()


# ========================================================================
# DATABASE (Module 8 - async SQLAlchemy)
# ========================================================================

engine = create_async_engine(settings.DATABASE_URL, echo=False, future=True)
AsyncSessionLocal = async_sessionmaker(bind=engine, class_=AsyncSession, expire_on_commit=False)


class Base(DeclarativeBase):
    pass


async def init_db() -> None:
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)


async def get_db():
    async with AsyncSessionLocal() as session:
        try:
            yield session
        finally:
            await session.close()


# ========================================================================
# MODELS
# ========================================================================

def _utcnow() -> datetime.datetime:
    return datetime.datetime.utcnow()


class ProcessingStatus(str, enum.Enum):
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"


class User(Base):
    __tablename__ = "users"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, index=True)
    email: Mapped[str] = mapped_column(String(255), unique=True, index=True, nullable=False)
    hashed_password: Mapped[str] = mapped_column(String(255), nullable=False)
    created_at: Mapped[datetime.datetime] = mapped_column(DateTime, default=_utcnow)

    documents: Mapped[list["Document"]] = relationship(
        "Document", back_populates="owner", cascade="all, delete-orphan"
    )


class Document(Base):
    __tablename__ = "documents"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, index=True)
    owner_id: Mapped[int] = mapped_column(ForeignKey("users.id"), nullable=False)

    filename: Mapped[str] = mapped_column(String(255), nullable=False)
    file_path: Mapped[str] = mapped_column(String(512), nullable=False)
    content_type: Mapped[str] = mapped_column(String(100), nullable=False)
    file_size_bytes: Mapped[int] = mapped_column(Integer, default=0)
    uploaded_at: Mapped[datetime.datetime] = mapped_column(DateTime, default=_utcnow)

    extracted_text: Mapped[str] = mapped_column(Text, default="")
    processing_status: Mapped[ProcessingStatus] = mapped_column(
        Enum(ProcessingStatus), default=ProcessingStatus.PENDING
    )

    owner: Mapped["User"] = relationship("User", back_populates="documents")
    analysis: Mapped["AIAnalysis"] = relationship(
        "AIAnalysis", back_populates="document", uselist=False, cascade="all, delete-orphan"
    )
    questions: Mapped[list["QARecord"]] = relationship(
        "QARecord", back_populates="document", cascade="all, delete-orphan"
    )


class AIAnalysis(Base):
    __tablename__ = "ai_analyses"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, index=True)
    document_id: Mapped[int] = mapped_column(ForeignKey("documents.id"), unique=True, nullable=False)

    summary: Mapped[str] = mapped_column(Text, default="")
    key_points: Mapped[str] = mapped_column(Text, default="")
    important_topics: Mapped[str] = mapped_column(Text, default="")
    created_at: Mapped[datetime.datetime] = mapped_column(DateTime, default=_utcnow)

    document: Mapped["Document"] = relationship("Document", back_populates="analysis")


class QARecord(Base):
    __tablename__ = "qa_records"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, index=True)
    document_id: Mapped[int] = mapped_column(ForeignKey("documents.id"), nullable=False)

    question: Mapped[str] = mapped_column(Text, nullable=False)
    answer: Mapped[str] = mapped_column(Text, nullable=False)
    created_at: Mapped[datetime.datetime] = mapped_column(DateTime, default=_utcnow)

    document: Mapped["Document"] = relationship("Document", back_populates="questions")


# ========================================================================
# SCHEMAS
# ========================================================================

class UserRegister(BaseModel):
    email: EmailStr
    password: str = Field(min_length=6)


class UserLogin(BaseModel):
    email: EmailStr
    password: str


class UserOut(BaseModel):
    model_config = ConfigDict(from_attributes=True)
    id: int
    email: EmailStr
    created_at: datetime.datetime


class Token(BaseModel):
    access_token: str
    token_type: str = "bearer"


class DocumentOut(BaseModel):
    model_config = ConfigDict(from_attributes=True)
    id: int
    filename: str
    content_type: str
    file_size_bytes: int
    uploaded_at: datetime.datetime
    processing_status: str


class DocumentDetailOut(DocumentOut):
    extracted_text: str


class DocumentUploadResponse(BaseModel):
    document: DocumentOut
    message: str = "File uploaded. AI analysis is running in the background."


class ManualTextInput(BaseModel):
    title: str = Field(min_length=1, max_length=255)
    text: str = Field(min_length=1)


class AIAnalysisOut(BaseModel):
    model_config = ConfigDict(from_attributes=True)
    document_id: int
    summary: str
    key_points: list[str]
    important_topics: list[str]


class QuestionRequest(BaseModel):
    question: str = Field(min_length=1)


class QAOut(BaseModel):
    model_config = ConfigDict(from_attributes=True)
    id: int
    document_id: int
    question: str
    answer: str
    created_at: datetime.datetime


# ========================================================================
# AUTH (Module 1) - password hashing + JWT
# ========================================================================

oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/auth/login")


def hash_password(password: str) -> str:
    password_bytes = password.encode("utf-8")[:72]
    return bcrypt.hashpw(password_bytes, bcrypt.gensalt()).decode("utf-8")


def verify_password(plain_password: str, hashed_password: str) -> bool:
    password_bytes = plain_password.encode("utf-8")[:72]
    try:
        return bcrypt.checkpw(password_bytes, hashed_password.encode("utf-8"))
    except ValueError:
        return False


def create_access_token(subject: str, expires_minutes: int | None = None) -> str:
    expire_minutes = expires_minutes or settings.ACCESS_TOKEN_EXPIRE_MINUTES
    expire = datetime.datetime.utcnow() + datetime.timedelta(minutes=expire_minutes)
    payload = {"sub": subject, "exp": expire}
    return jwt.encode(payload, settings.SECRET_KEY, algorithm=settings.ALGORITHM)


def decode_access_token(token: str) -> str:
    try:
        payload = jwt.decode(token, settings.SECRET_KEY, algorithms=[settings.ALGORITHM])
        subject: str = payload.get("sub")
        if subject is None:
            raise JWTError("Missing subject claim")
        return subject
    except JWTError as exc:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Could not validate credentials",
            headers={"WWW-Authenticate": "Bearer"},
        ) from exc


async def get_current_user(
    token: str = Depends(oauth2_scheme), db: AsyncSession = Depends(get_db)
) -> User:
    email = decode_access_token(token)
    result = await db.execute(select(User).where(User.email == email))
    user = result.scalar_one_or_none()
    if user is None:
        logger.warning("Auth failed: token valid but user %s no longer exists", email)
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="User not found",
            headers={"WWW-Authenticate": "Bearer"},
        )
    return user


# ========================================================================
# FILE PROCESSING (Module 3) - text extraction
# ========================================================================

SUPPORTED_CONTENT_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}


def _extract_from_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _extract_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text).strip()


def _extract_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read().strip()


def extract_text(file_path: str, file_type: str) -> str:
    try:
        if file_type == "pdf":
            return _extract_from_pdf(file_path)
        elif file_type == "docx":
            return _extract_from_docx(file_path)
        elif file_type == "txt":
            return _extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    except Exception:
        logger.exception("Text extraction failed for %s (%s)", file_path, file_type)
        raise


def resolve_file_type(content_type: str, filename: str) -> str:
    if content_type in SUPPORTED_CONTENT_TYPES:
        return SUPPORTED_CONTENT_TYPES[content_type]

    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        return "pdf"
    if ext == ".docx":
        return "docx"
    if ext == ".txt":
        return "txt"

    raise ValueError(
        f"Unsupported file: content_type={content_type}, filename={filename}. "
        "Only PDF, DOCX and TXT are supported."
    )


# ========================================================================
# PROMPT TEMPLATES (Module 10)
# ========================================================================

SUMMARY_MARKER = "SUMMARY:"
KEY_POINTS_MARKER = "KEY_POINTS:"
TOPICS_MARKER = "IMPORTANT_TOPICS:"

DOCUMENT_ANALYSIS_TEMPLATE_STR = """You are an assistant that analyzes documents.

Read the document below and produce exactly three sections, using these
exact section headers (all caps, followed by a colon), with nothing else
before, between, or after them:

{summary_marker}
<a concise 3-6 sentence summary of the document>

{key_points_marker}
- <key point 1>
- <key point 2>
- <key point 3>
(as many bullet points as are genuinely useful, one per line, each starting with "- ")

{topics_marker}
- <important topic 1>
- <important topic 2>
(one topic per line, each starting with "- ")

Document:
\"\"\"
{text}
\"\"\"
"""

document_analysis_prompt = PromptTemplate(
    input_variables=["text"],
    partial_variables={
        "summary_marker": SUMMARY_MARKER,
        "key_points_marker": KEY_POINTS_MARKER,
        "topics_marker": TOPICS_MARKER,
    },
    template=DOCUMENT_ANALYSIS_TEMPLATE_STR,
)


def build_document_analysis_prompt(text: str) -> str:
    return document_analysis_prompt.format(text=text)


QUESTION_ANSWER_TEMPLATE_STR = """You are a helpful assistant answering questions about a specific document.

Only use information contained in the document below. If the answer is
not present in the document, say clearly that the document does not
contain that information - do not make anything up.

Document:
\"\"\"
{text}
\"\"\"

Question: {question}

Answer:"""

question_answer_prompt = PromptTemplate(
    input_variables=["text", "question"], template=QUESTION_ANSWER_TEMPLATE_STR
)


def build_question_answer_prompt(text: str, question: str) -> str:
    return question_answer_prompt.format(text=text, question=question)


# ========================================================================
# AI SERVICE (Module 4, 9, 11) - Gemini + LangChain + output parsing
# ========================================================================

@dataclass
class DocumentAnalysisResult:
    summary: str
    key_points: list[str]
    important_topics: list[str]


class GeminiUnavailableError(RuntimeError):
    """Raised when no Gemini API key is configured."""


def _require_api_key() -> None:
    if not settings.GEMINI_API_KEY:
        raise GeminiUnavailableError(
            "GEMINI_API_KEY is not set. Add it to your .env file to enable AI features."
        )


def _get_langchain_llm():
    from langchain_google_genai import ChatGoogleGenerativeAI

    return ChatGoogleGenerativeAI(
        model=settings.GEMINI_MODEL, google_api_key=settings.GEMINI_API_KEY, temperature=0.3
    )


def _call_gemini_via_langchain(rendered_prompt: str) -> str:
    """Module 11: run the rendered prompt through LangChain's LCEL pipeline."""
    llm = _get_langchain_llm()
    response = llm.invoke(rendered_prompt)
    return response.content


def _call_gemini_direct(rendered_prompt: str) -> str:
    """Fallback: call google-generativeai directly (no LangChain)."""
    import google.generativeai as genai

    genai.configure(api_key=settings.GEMINI_API_KEY)
    model = genai.GenerativeModel(settings.GEMINI_MODEL)
    response = model.generate_content(rendered_prompt)
    return response.text


def _call_gemini(rendered_prompt: str) -> str:
    _require_api_key()
    try:
        return _call_gemini_via_langchain(rendered_prompt)
    except Exception as exc:  # noqa: BLE001
        logger.warning("LangChain call failed (%s); falling back to direct SDK call", exc)
        return _call_gemini_direct(rendered_prompt)


def _extract_bullet_lines(block: str) -> list[str]:
    lines = []
    for line in block.strip().splitlines():
        line = line.strip()
        if not line:
            continue
        line = re.sub(r"^[-*\u2022]\s*", "", line)
        if line:
            lines.append(line)
    return lines


def _parse_analysis_response(raw_text: str) -> DocumentAnalysisResult:
    """Module 9: split Gemini's raw response into summary / key_points / topics."""
    pattern = re.compile(
        rf"{re.escape(SUMMARY_MARKER)}(?P<summary>.*?)"
        rf"{re.escape(KEY_POINTS_MARKER)}(?P<key_points>.*?)"
        rf"{re.escape(TOPICS_MARKER)}(?P<topics>.*)",
        re.DOTALL,
    )
    match = pattern.search(raw_text)

    if not match:
        logger.warning("Could not parse structured sections from Gemini response; using raw text as summary")
        return DocumentAnalysisResult(summary=raw_text.strip(), key_points=[], important_topics=[])

    return DocumentAnalysisResult(
        summary=match.group("summary").strip(),
        key_points=_extract_bullet_lines(match.group("key_points")),
        important_topics=_extract_bullet_lines(match.group("topics")),
    )


async def analyze_document(text: str) -> DocumentAnalysisResult:
    """Module 4/9/11: Summary + Key Points + Important Topics for a document."""
    prompt = build_document_analysis_prompt(text=text)
    raw_response = await asyncio.to_thread(_call_gemini, prompt)
    return _parse_analysis_response(raw_response)


async def answer_question(text: str, question: str) -> str:
    """Module 5: answer a user's question using only the document text."""
    prompt = build_question_answer_prompt(text=text, question=question)
    raw_response = await asyncio.to_thread(_call_gemini, prompt)
    return raw_response.strip()


# ========================================================================
# BACKGROUND TASKS (Module 7)
# ========================================================================

async def run_document_analysis(document_id: int) -> None:
    """
    Runs after the upload request has already responded (Module 7).
    Opens its own DB session since the request's session is gone by then.
    """
    async with AsyncSessionLocal() as db:
        result = await db.execute(select(Document).where(Document.id == document_id))
        document = result.scalar_one_or_none()
        if document is None:
            logger.error("Background analysis: document %s not found", document_id)
            return

        document.processing_status = ProcessingStatus.PROCESSING
        await db.commit()

        try:
            logger.info("AI request started for document_id=%s", document_id)
            analysis_result = await analyze_document(document.extracted_text)

            existing = await db.execute(select(AIAnalysis).where(AIAnalysis.document_id == document_id))
            ai_row = existing.scalar_one_or_none()
            if ai_row is None:
                ai_row = AIAnalysis(document_id=document_id)
                db.add(ai_row)

            ai_row.summary = analysis_result.summary
            ai_row.key_points = "\n".join(analysis_result.key_points)
            ai_row.important_topics = "\n".join(analysis_result.important_topics)

            document.processing_status = ProcessingStatus.COMPLETED
            await db.commit()
            logger.info("AI request completed for document_id=%s", document_id)

        except Exception:
            logger.exception("AI request failed for document_id=%s", document_id)
            document.processing_status = ProcessingStatus.FAILED
            await db.commit()


# ========================================================================
# FASTAPI APP + MIDDLEWARE (Module 6 - log every API call)
# ========================================================================

app = FastAPI(
    title="AI-Powered Knowledge Base API",
    description=(
        "Upload PDF/DOCX/TXT documents, get AI-generated summaries, key points and "
        "topics, and ask questions about your documents using Gemini."
    ),
    version="1.0.0",
)


@app.on_event("startup")
async def on_startup():
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    os.makedirs(settings.LOG_DIR, exist_ok=True)
    await init_db()
    logger.info("Application startup complete. Tables ensured, folders ready.")


@app.middleware("http")
async def log_requests(request: Request, call_next):
    start = time.perf_counter()
    try:
        response = await call_next(request)
    except Exception:
        logger.exception("Unhandled error on %s %s", request.method, request.url.path)
        raise
    duration_ms = (time.perf_counter() - start) * 1000
    logger.info(
        "API call: %s %s -> %s (%.1fms)",
        request.method,
        request.url.path,
        response.status_code,
        duration_ms,
    )
    return response


@app.exception_handler(Exception)
async def unhandled_exception_handler(request: Request, exc: Exception):
    logger.error("Unhandled exception on %s %s: %s", request.method, request.url.path, exc)
    return JSONResponse(status_code=500, content={"detail": "Internal server error"})


# ========================================================================
# ROUTES - Module 1: Authentication
# ========================================================================

@app.post("/auth/register", response_model=UserOut, status_code=status.HTTP_201_CREATED, tags=["Authentication"])
async def register_user(payload: UserRegister, db: AsyncSession = Depends(get_db)):
    existing = await db.execute(select(User).where(User.email == payload.email))
    if existing.scalar_one_or_none() is not None:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Email already registered")

    user = User(email=payload.email, hashed_password=hash_password(payload.password))
    db.add(user)
    await db.commit()
    await db.refresh(user)

    logger.info("New user registered: %s", user.email)
    return user


@app.post("/auth/login", response_model=Token, tags=["Authentication"])
async def login_user(payload: UserLogin, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(User).where(User.email == payload.email))
    user = result.scalar_one_or_none()

    if user is None or not verify_password(payload.password, user.hashed_password):
        logger.warning("Failed login attempt for email=%s", payload.email)
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid email or password")

    token = create_access_token(subject=user.email)
    logger.info("User login: %s", user.email)
    return Token(access_token=token)


# ========================================================================
# ROUTES - Module 2/3/7: Documents (upload, list, detail, analysis)
# ========================================================================

MAX_UPLOAD_BYTES = settings.MAX_UPLOAD_MB * 1024 * 1024


async def _get_owned_document(document_id: int, db: AsyncSession, current_user: User) -> Document:
    result = await db.execute(select(Document).where(Document.id == document_id))
    document = result.scalar_one_or_none()
    if document is None or document.owner_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")
    return document


@app.post(
    "/documents/upload",
    response_model=DocumentUploadResponse,
    status_code=status.HTTP_201_CREATED,
    tags=["Documents"],
)
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Upload a PDF/DOCX/TXT file: save to disk, extract text immediately,
    then schedule AI summarization as a BackgroundTask so the response
    returns right away (Module 7).
    """
    try:
        file_type = resolve_file_type(file.content_type or "", file.filename or "")
    except ValueError as exc:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc)) from exc

    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    safe_name = f"{uuid.uuid4().hex}_{file.filename}"
    disk_path = os.path.join(settings.UPLOAD_DIR, safe_name)

    contents = await file.read()
    if len(contents) > MAX_UPLOAD_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File exceeds {settings.MAX_UPLOAD_MB}MB limit",
        )

    with open(disk_path, "wb") as f:
        f.write(contents)

    logger.info("File uploaded by user=%s: %s (%s bytes)", current_user.email, file.filename, len(contents))

    try:
        extracted_text = await asyncio.to_thread(extract_text, disk_path, file_type)
    except Exception as exc:
        logger.error("Text extraction failed for %s: %s", file.filename, exc)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not extract text from the uploaded file",
        ) from exc

    document = Document(
        owner_id=current_user.id,
        filename=file.filename or safe_name,
        file_path=disk_path,
        content_type=file.content_type or "application/octet-stream",
        file_size_bytes=len(contents),
        extracted_text=extracted_text,
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)

    background_tasks.add_task(run_document_analysis, document.id)

    return DocumentUploadResponse(document=DocumentOut.model_validate(document))


@app.post(
    "/documents/manual",
    response_model=DocumentUploadResponse,
    status_code=status.HTTP_201_CREATED,
    tags=["Documents"],
)
async def add_manual_text(
    payload: ManualTextInput,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Create a "document" from text the user types/pastes directly instead
    of uploading a file. The text is saved to disk as a .txt file so it
    behaves exactly like an uploaded document (can be deleted, re-read,
    etc.), and AI summarization runs in the background just like uploads.
    """
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    safe_filename = f"{payload.title.strip().replace(' ', '_')}.txt"
    safe_name = f"{uuid.uuid4().hex}_{safe_filename}"
    disk_path = os.path.join(settings.UPLOAD_DIR, safe_name)

    text_bytes = payload.text.encode("utf-8")
    with open(disk_path, "wb") as f:
        f.write(text_bytes)

    logger.info("Manual text input added by user=%s: %s (%s bytes)", current_user.email, payload.title, len(text_bytes))

    document = Document(
        owner_id=current_user.id,
        filename=safe_filename,
        file_path=disk_path,
        content_type="text/plain",
        file_size_bytes=len(text_bytes),
        extracted_text=payload.text.strip(),
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)

    background_tasks.add_task(run_document_analysis, document.id)

    return DocumentUploadResponse(document=DocumentOut.model_validate(document))


@app.get("/documents", response_model=list[DocumentOut], tags=["Documents"])
async def list_documents(
    q: str | None = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    List the current user's documents. Optional `q` performs a search:
    if `q` is numeric it matches an exact document id, otherwise it does
    a case-insensitive substring match against the filename.
    """
    query = select(Document).where(Document.owner_id == current_user.id)

    if q:
        q = q.strip()
        if q.isdigit():
            query = query.where(Document.id == int(q))
        else:
            query = query.where(Document.filename.ilike(f"%{q}%"))

    result = await db.execute(query)
    return result.scalars().all()


@app.get("/documents/{document_id}", response_model=DocumentDetailOut, tags=["Documents"])
async def get_document(
    document_id: int, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)
):
    return await _get_owned_document(document_id, db, current_user)


@app.get("/documents/{document_id}/analysis", response_model=AIAnalysisOut, tags=["Documents"])
async def get_document_analysis(
    document_id: int, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)
):
    document = await _get_owned_document(document_id, db, current_user)

    result = await db.execute(select(AIAnalysis).where(AIAnalysis.document_id == document.id))
    analysis = result.scalar_one_or_none()

    if analysis is None:
        raise HTTPException(
            status_code=status.HTTP_202_ACCEPTED,
            detail=f"AI analysis status: {document.processing_status.value}. Not ready yet.",
        )

    return AIAnalysisOut(
        document_id=document.id,
        summary=analysis.summary,
        key_points=[p for p in analysis.key_points.split("\n") if p],
        important_topics=[t for t in analysis.important_topics.split("\n") if t],
    )


def _build_export_docx(document: Document, analysis: "AIAnalysis | None", qa_records: list["QARecord"]) -> BytesIO:
    """Build a Word document containing a document's AI analysis and Q&A history."""
    docx_doc = DocxDocument()

    docx_doc.add_heading(document.filename, level=0)
    docx_doc.add_paragraph(f"Uploaded: {document.uploaded_at.strftime('%Y-%m-%d %H:%M')}")
    docx_doc.add_paragraph(f"Status: {document.processing_status.value}")

    if analysis is not None:
        docx_doc.add_heading("Summary", level=1)
        docx_doc.add_paragraph(analysis.summary or "—")

        docx_doc.add_heading("Key Points", level=1)
        key_points = [p for p in analysis.key_points.split("\n") if p]
        if key_points:
            for point in key_points:
                docx_doc.add_paragraph(point, style="List Bullet")
        else:
            docx_doc.add_paragraph("—")

        docx_doc.add_heading("Important Topics", level=1)
        topics = [t for t in analysis.important_topics.split("\n") if t]
        if topics:
            for topic in topics:
                docx_doc.add_paragraph(topic, style="List Bullet")
        else:
            docx_doc.add_paragraph("—")
    else:
        docx_doc.add_heading("AI Analysis", level=1)
        docx_doc.add_paragraph("AI analysis is not available for this document yet.")

    if qa_records:
        docx_doc.add_heading("Questions & Answers", level=1)
        for qa in qa_records:
            docx_doc.add_heading(qa.question, level=2)
            docx_doc.add_paragraph(qa.answer)

    buffer = BytesIO()
    docx_doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.get("/documents/{document_id}/export", tags=["Documents"])
async def export_document_word(
    document_id: int, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)
):
    """Export a document's AI summary, key points, topics, and Q&A history as a .docx file."""
    document = await _get_owned_document(document_id, db, current_user)

    analysis_result = await db.execute(select(AIAnalysis).where(AIAnalysis.document_id == document.id))
    analysis = analysis_result.scalar_one_or_none()

    qa_result = await db.execute(select(QARecord).where(QARecord.document_id == document.id))
    qa_records = qa_result.scalars().all()

    buffer = _build_export_docx(document, analysis, qa_records)
    safe_name = re.sub(r"[^A-Za-z0-9_.-]", "_", document.filename.rsplit(".", 1)[0])

    logger.info("Document %s exported to Word by user=%s", document_id, current_user.email)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}_export.docx"'},
    )


# ========================================================================
# ROUTES - Module 5: Question Answering
# ========================================================================

@app.post(
    "/documents/{document_id}/ask",
    response_model=QAOut,
    status_code=status.HTTP_201_CREATED,
    tags=["Question Answering"],
)
async def ask_question(
    document_id: int,
    payload: QuestionRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    document = await _get_owned_document(document_id, db, current_user)

    if not document.extracted_text:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="This document has no extracted text to answer from.",
        )

    logger.info("AI request (QA) for document_id=%s by user=%s", document_id, current_user.email)
    try:
        answer = await answer_question(document.extracted_text, payload.question)
    except GeminiUnavailableError as exc:
        raise HTTPException(status_code=status.HTTP_503_SERVICE_UNAVAILABLE, detail=str(exc)) from exc
    except Exception as exc:
        logger.exception("QA failed for document_id=%s", document_id)
        raise HTTPException(status_code=status.HTTP_502_BAD_GATEWAY, detail="AI provider request failed") from exc

    qa_record = QARecord(document_id=document.id, question=payload.question, answer=answer)
    db.add(qa_record)
    await db.commit()
    await db.refresh(qa_record)

    return qa_record


@app.get("/documents/{document_id}/questions", response_model=list[QAOut], tags=["Question Answering"])
async def list_questions(
    document_id: int, db: AsyncSession = Depends(get_db), current_user: User = Depends(get_current_user)
):
    document = await _get_owned_document(document_id, db, current_user)
    result = await db.execute(select(QARecord).where(QARecord.document_id == document.id))
    return result.scalars().all()

@app.delete("/documents/{document_id}", tags=["Documents"])
async def delete_document(
    document_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    document = await _get_owned_document(document_id, db, current_user)

    # Delete the file from disk
    if os.path.exists(document.file_path):
        os.remove(document.file_path)

    # Delete the database record
    await db.delete(document)
    await db.commit()

    logger.info("Document %s deleted by %s", document.filename, current_user.email)

    return {"message": "Document deleted successfully."}


# ========================================================================
# HEALTH CHECK
# ========================================================================

@app.get("/health", tags=["Health"])
async def health_check():
    return {"status": "ok"}